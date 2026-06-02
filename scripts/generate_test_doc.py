from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(r"D:\AI Demo\菁选")
MARKDOWN_PATH = PROJECT_ROOT / "docs" / "功能测试文档.md"
OUTPUT_PATH = PROJECT_ROOT / "docs" / "菁选校园作品展示平台功能测试文档_v2.docx"


@dataclass
class CaseRow:
    module: str
    subfeature: str
    precondition: str
    expected: str
    automation: str


@dataclass
class CaseGroup:
    domain: str
    cases: list[CaseRow]


def parse_markdown(md_text: str) -> dict:
    lines = md_text.splitlines()
    meta = {}
    current_h2 = None
    current_h3 = None
    overview = []
    groups: list[CaseGroup] = []
    entry_conditions = []
    exit_conditions = []
    env_rows = []
    release_checks = []
    integrity_rows = []

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if stripped.startswith("> "):
            body = stripped[2:]
            if "：" in body:
                key, value = body.split("：", 1)
                meta[key.strip()] = value.strip()
        elif stripped.startswith("## "):
            current_h2 = stripped[3:].strip()
            current_h3 = None
        elif stripped.startswith("### "):
            current_h3 = stripped[4:].strip()
        elif stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            table = parse_table(table_lines)
            if current_h2 == "1. 功能模块概述":
                overview = table
            elif current_h2 == "2. 功能测试矩阵" and current_h3:
                groups.append(
                    CaseGroup(
                        domain=current_h3,
                        cases=[
                            CaseRow(
                                module=row.get("功能", "").strip(),
                                subfeature=row.get("子功能", "").strip(),
                                precondition=row.get("前置条件", "").strip(),
                                expected=row.get("预期结果", "").strip(),
                                automation=row.get("自动化", "").strip(),
                            )
                            for row in table
                        ],
                    )
                )
            elif current_h2 == "4. 功能测试环境":
                env_rows = table
            elif current_h2 == "5. 功能测试 Checklist" and current_h3 == "5.2 功能完整性检查":
                integrity_rows = table
            continue
        elif stripped.startswith("- ["):
            if current_h2 == "5. 功能测试 Checklist" and current_h3 == "5.1 发布前检查清单":
                release_checks.append(stripped)
        elif stripped.startswith("- "):
            bullet = stripped[2:].strip()
            if current_h2 == "3. 功能测试准入/准出" and current_h3 == "3.1 准入条件":
                entry_conditions.append(bullet)
            elif current_h2 == "3. 功能测试准入/准出" and current_h3 == "3.2 准出条件":
                exit_conditions.append(bullet)
        i += 1

    return {
        "meta": meta,
        "overview": overview,
        "groups": groups,
        "entry_conditions": entry_conditions,
        "exit_conditions": exit_conditions,
        "environments": env_rows,
        "release_checks": release_checks,
        "integrity_rows": integrity_rows,
    }


def parse_table(lines: list[str]) -> list[dict]:
    rows = []
    parsed = []
    for line in lines:
        parsed.append([chunk.strip() for chunk in line.strip().strip("|").split("|")])
    if len(parsed) < 2:
        return rows
    headers = parsed[0]
    for row in parsed[2:]:
        if len(row) < len(headers):
            row += [""] * (len(headers) - len(row))
        rows.append(dict(zip(headers, row)))
    return rows


def set_run_font(run, size_pt=10.5, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "SimSun")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "SimSun")
    run.font.size = Pt(size_pt)
    run.bold = bold


def style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.2


def add_text(doc: Document, text: str, size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6):
    p = doc.add_paragraph()
    style_paragraph(p, align=align, space_before=before, space_after=after)
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    return add_text(
        doc,
        text,
        size=14 if level == 1 else 12,
        bold=True,
        before=10 if level == 1 else 6,
        after=6,
    )


def set_page_layout(section, landscape=False):
    portrait_width = Cm(21.0)
    portrait_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.8 if landscape else 2.2)
    section.right_margin = Cm(1.8 if landscape else 2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = portrait_height
        section.page_height = portrait_width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = portrait_width
        section.page_height = portrait_height


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.5, bold=False, shading=None):
    cell.text = ""
    p = cell.paragraphs[0]
    style_paragraph(p, align=align, space_before=0, space_after=0)
    lines = str(text).split("\n")
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, size_pt=size, bold=bold)
        if idx < len(lines) - 1:
            run.add_break()
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if shading:
        set_cell_shading(cell, shading)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")


def build_summary_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths_cm: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Cm(col_widths_cm[idx])
        set_cell_text(cell, header, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, bold=True, shading="D9E2F3")
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].width = Cm(col_widths_cm[idx])
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], value, align=align, size=10)
    return table


def normalize_automation(value: str) -> str:
    if "✅" in value:
        return "自动化"
    if "❌" in value:
        return "手工"
    return value or ""


def build_step_text(case: CaseRow) -> str:
    action = case.subfeature or case.module
    if case.precondition and case.precondition not in {"—", "-"}:
        return f"1. 进入{case.module}\n2. 执行“{action}”\n3. 校验返回结果"
    return f"1. 进入{case.module}\n2. 执行“{action}”\n3. 观察系统反馈"


def add_case_meta(doc: Document, meta: dict, group_name: str, module_name: str, case_count: int):
    meta_rows = [
        ["项目名称", meta.get("项目", "")],
        ["模块", group_name],
        ["功能点", module_name],
        ["版本", meta.get("版本", "")],
        ["更新日期", meta.get("更新", "")],
        ["用例数量", str(case_count)],
    ]
    build_summary_table(doc, ["字段", "内容"], meta_rows, [3.2, 17.8])


def build_case_table(doc: Document, cases: list[CaseRow], start_index: int) -> int:
    headers = ["编号", "功能", "子功能/场景", "前置条件", "操作步骤", "预期结果", "执行方式"]
    widths = [1.6, 2.1, 3.2, 3.4, 5.2, 5.2, 1.6]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].width = Cm(widths[idx])
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=10,
            bold=True,
            shading="BFBFBF",
        )

    for case in cases:
        row = table.add_row()
        values = [
            f"FT-{start_index:03d}",
            case.module,
            case.subfeature or case.module,
            case.precondition or "—",
            build_step_text(case),
            case.expected or "—",
            normalize_automation(case.automation),
        ]
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
        ]
        for idx, value in enumerate(values):
            row.cells[idx].width = Cm(widths[idx])
            set_cell_text(row.cells[idx], value, align=aligns[idx], size=9.5)
        start_index += 1
    return start_index


def add_case_groups(doc: Document, parsed: dict):
    landscape_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_layout(landscape_section, landscape=True)
    add_heading(doc, "2. 功能测试用例明细", level=1)

    case_index = 1
    for group in parsed["groups"]:
        add_heading(doc, group.domain, level=2)
        grouped: dict[str, list[CaseRow]] = {}
        for case in group.cases:
            grouped.setdefault(case.module, []).append(case)
        for module_name, cases in grouped.items():
            add_text(doc, module_name, size=11, bold=True, before=4, after=4)
            add_case_meta(doc, parsed["meta"], group.domain, module_name, len(cases))
            doc.add_paragraph()
            case_index = build_case_table(doc, cases, case_index)
            doc.add_paragraph()

    portrait_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page_layout(portrait_section, landscape=False)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.space_after = Pt(4)
        run1 = p.add_run("• ")
        set_run_font(run1, size_pt=10.5, bold=True)
        run2 = p.add_run(item)
        set_run_font(run2, size_pt=10.5)


def main():
    parsed = parse_markdown(MARKDOWN_PATH.read_text(encoding="utf-8"))
    doc = Document()
    set_page_layout(doc.sections[0], landscape=False)

    add_text(doc, parsed["meta"].get("项目", "功能测试文档"), size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=12)
    add_text(doc, "功能测试文档", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    build_summary_table(
        doc,
        ["字段", "内容"],
        [
            ["项目名称", parsed["meta"].get("项目", "")],
            ["版本", parsed["meta"].get("版本", "")],
            ["更新日期", parsed["meta"].get("更新", "")],
            ["文档来源", "依据《功能测试文档.md》整理生成"],
        ],
        [3.2, 12.3],
    )
    doc.add_paragraph()

    add_heading(doc, "1. 功能模块概述", level=1)
    build_summary_table(
        doc,
        ["模块", "角色", "核心功能"],
        [[row["模块"], row["角色"], row["核心功能"]] for row in parsed["overview"]],
        [2.5, 2.8, 10.2],
    )

    add_case_groups(doc, parsed)

    add_heading(doc, "3. 准入与准出条件", level=1)
    add_heading(doc, "3.1 准入条件", level=2)
    add_bullets(doc, parsed["entry_conditions"])
    add_heading(doc, "3.2 准出条件", level=2)
    add_bullets(doc, parsed["exit_conditions"])

    add_heading(doc, "4. 功能测试环境", level=1)
    build_summary_table(
        doc,
        ["环境", "用途", "数据库", "外部依赖"],
        [[row["环境"], row["用途"], row["数据库"], row["外部依赖"]] for row in parsed["environments"]],
        [2.5, 3.0, 3.2, 6.8],
    )

    add_heading(doc, "5. 发布前检查清单", level=1)
    checklist_rows = []
    for item in parsed["release_checks"]:
        matched = re.match(r"- \[(.)\] (.+)", item)
        if matched:
            checklist_rows.append(["已完成" if matched.group(1).lower() == "x" else "未完成", matched.group(2)])
    build_summary_table(doc, ["状态", "检查项"], checklist_rows, [2.5, 13.0])

    add_heading(doc, "6. 功能完整性检查", level=1)
    build_summary_table(
        doc,
        ["检查项", "验证方式", "通过标准"],
        [[row["检查项"], row["验证方式"], row["通过标准"]] for row in parsed["integrity_rows"]],
        [4.5, 3.0, 8.0],
    )

    doc.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
