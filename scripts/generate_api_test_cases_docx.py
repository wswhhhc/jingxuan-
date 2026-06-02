from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


AUTHOR = "Codex"
DATE_TEXT = "2026-05-25"
CASE_PREFIX = "JX-API"


@dataclass
class ApiCase:
    module_title: str
    api_title: str
    api_path: str
    method: str
    coverage: str
    summary: str = ""
    request_body: str = ""
    notes: List[str] = field(default_factory=list)
    tables: List[dict] = field(default_factory=list)


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False, font_size: int = 10):
    cell.text = ""
    paragraphs = text.split("\n") if text else [""]
    for idx, line in enumerate(paragraphs):
        para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        run = para.add_run(line)
        run.bold = bold
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(font_size)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.15
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def set_table_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def parse_markdown(md_text: str) -> List[ApiCase]:
    lines = md_text.splitlines()
    modules: List[ApiCase] = []
    current_module = ""
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("## ") and re.match(r"## \d+\.", line):
            current_module = line[3:].strip()
            i += 1
            continue
        if line.startswith("### ") and re.match(r"### \d+\.\d+", line):
            title = line[4:].strip()
            method_match = re.search(r"(GET|POST|PUT|DELETE|PATCH)\s+`([^`]+)`", title)
            coverage_match = re.search(r"`(\[[^`]+\])`", title)
            api = ApiCase(
                module_title=current_module,
                api_title=title,
                api_path=method_match.group(2) if method_match else title,
                method=method_match.group(1) if method_match else "",
                coverage=coverage_match.group(1) if coverage_match else "",
            )
            i += 1
            block: List[str] = []
            while i < len(lines):
                next_line = lines[i].rstrip()
                if next_line.startswith("### ") and re.match(r"### \d+\.\d+", next_line):
                    break
                if next_line.startswith("## ") and re.match(r"## \d+\.", next_line):
                    break
                block.append(next_line)
                i += 1
            consume_block_into_api(block, api)
            modules.append(api)
            continue
        i += 1
    return modules


def consume_block_into_api(block: List[str], api: ApiCase) -> None:
    i = 0
    active_label = ""
    prose_bucket: List[str] = []
    while i < len(block):
        line = block[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("**") and stripped.endswith("：**"):
            active_label = stripped.strip("*").rstrip("：:")
            i += 1
            continue
        if stripped.startswith("**请求体：**"):
            code_lines = []
            i += 1
            while i < len(block) and not block[i].strip().startswith("```"):
                i += 1
            if i < len(block) and block[i].strip().startswith("```"):
                i += 1
                while i < len(block) and not block[i].strip().startswith("```"):
                    code_lines.append(block[i])
                    i += 1
                api.request_body = "\n".join(code_lines).strip()
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = [stripped]
            i += 1
            while i < len(block) and block[i].strip().startswith("|"):
                table_lines.append(block[i].strip())
                i += 1
            table = parse_md_table(table_lines)
            api.tables.append({"label": active_label or "用例", "table": table})
            continue
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(block) and not block[i].strip().startswith("```"):
                code_lines.append(block[i])
                i += 1
            api.notes.append("代码示例：\n" + "\n".join(code_lines).strip())
            i += 1
            continue
        if stripped.startswith(">"):
            api.notes.append(stripped.lstrip("> ").strip())
            i += 1
            continue
        if stripped.startswith("- "):
            api.notes.append(stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("| API |"):
            i += 1
            continue
        prose_bucket.append(stripped)
        i += 1
    if prose_bucket:
        api.summary = "\n".join(prose_bucket[:4])


def parse_md_table(lines: List[str]) -> List[List[str]]:
    rows = []
    for idx, line in enumerate(lines):
        parts = [p.strip() for p in line.strip().strip("|").split("|")]
        if idx == 1 and all(set(p) <= {"-", ":"} for p in parts):
            continue
        rows.append(parts)
    return rows


def summarize_cases(table_rows: List[List[str]]) -> str:
    if not table_rows:
        return ""
    header = table_rows[0]
    body = table_rows[1:]
    chunks = []
    for row in body:
        pairs = []
        for h, v in zip(header, row):
            if v and h:
                pairs.append(f"{h}：{v}")
        chunks.append("；".join(pairs))
    return "\n".join(chunks)


def add_paragraph(doc: Document, text: str, style: str, *, bold=False, size=None, color=None):
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if run.font.name is not None or True:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return para


def build_case_table(doc: Document, api: ApiCase, index: int) -> None:
    table = doc.add_table(rows=7, cols=8)
    table.autofit = False
    widths = [Cm(2.0), Cm(2.4), Cm(1.7), Cm(1.7), Cm(1.6), Cm(1.7), Cm(1.7), Cm(3.1)]
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width

    set_table_borders(table)
    set_table_cell_margins(table)

    row0 = table.rows[0].cells
    set_cell_text(row0[0], "用例编号", bold=True, center=True)
    set_cell_text(row0[1], f"{CASE_PREFIX}-{index:04d}", center=True)
    set_cell_text(row0[2], "编写人", bold=True, center=True)
    set_cell_text(row0[3], AUTHOR, center=True)
    set_cell_text(row0[4], "编写时间", bold=True, center=True)
    merged_date = row0[5].merge(row0[7])
    set_cell_text(merged_date, DATE_TEXT, center=True)
    row0[5].text = ""
    row0[6].text = ""

    row1 = table.rows[1].cells
    set_cell_text(row1[0], "接口名称", bold=True, center=True)
    merged = row1[1].merge(row1[7])
    set_cell_text(merged, api.api_title)

    row2 = table.rows[2].cells
    set_cell_text(row2[0], "请求路径", bold=True, center=True)
    merged = row2[1].merge(row2[3])
    set_cell_text(merged, f"{api.method} {api.api_path}")
    set_cell_text(row2[4], "覆盖标记", bold=True, center=True)
    merged = row2[5].merge(row2[7])
    set_cell_text(merged, api.coverage or "未标注", center=True)

    row3 = table.rows[3].cells
    set_cell_text(row3[0], "接口说明", bold=True, center=True)
    merged = row3[1].merge(row3[7])
    summary = api.summary or "详见下方用例说明。"
    set_cell_text(merged, summary)

    row4 = table.rows[4].cells
    set_cell_text(row4[0], "请求体", bold=True, center=True)
    merged = row4[1].merge(row4[7])
    set_cell_text(merged, api.request_body or "无 / 按接口参数传递")

    row5 = table.rows[5].cells
    set_cell_text(row5[0], "测试要点", bold=True, center=True)
    merged = row5[1].merge(row5[7])
    notes = api.notes[:6]
    if not notes and api.tables:
        notes = [f"包含 {len(api.tables)} 组用例表，覆盖正向、异常或边界场景。"]
    set_cell_text(merged, "\n".join(notes) if notes else "覆盖鉴权、参数校验、业务规则与结果返回。")

    row6 = table.rows[6].cells
    set_cell_text(row6[0], "用例明细", bold=True, center=True)
    merged = row6[1].merge(row6[7])
    detail_parts = []
    for table_item in api.tables:
        detail_parts.append(f"{table_item['label']}：")
        detail_parts.append(summarize_cases(table_item["table"]))
    detail_text = "\n".join(part for part in detail_parts if part).strip()
    set_cell_text(merged, detail_text[:4000] if detail_text else "无")

    for ridx, row in enumerate(table.rows):
        for cidx, cell in enumerate(row.cells):
            if ridx == 0 or cidx == 0 or (ridx == 2 and cidx == 4):
                shade_cell(cell, "D9E2F3")


def build_document(template_path: Path, markdown_path: Path, output_path: Path) -> None:
    md_text = markdown_path.read_text(encoding="utf-8")
    apis = parse_markdown(md_text)
    doc = Document(str(template_path))
    clear_document_body(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("接口测试用例")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)
    title.paragraph_format.space_after = Pt(6)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("项目：菁选校园作品展示平台    版本：1.0.0    更新：2026-05-25")
    meta_run.font.name = "宋体"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    meta_run.font.size = Pt(10.5)
    meta.paragraph_format.space_after = Pt(12)

    add_paragraph(doc, "阅读指南", "Heading 2")
    for line in [
        "本文档依据《接口测试用例.md》整理，按角色模块拆分接口并沿用模板中的“标题 + 用例表”结构。",
        "返回约定：成功返回 Result<T>；业务异常返回 code=400；未认证返回 401；无权限返回 403；资源不存在返回 404。",
        "通用前置条件：认证接口需携带 Bearer token；测试前执行 test_data.sql 初始化数据；角色权限 ADMIN=3、TEACHER=2、STUDENT=1。",
    ]:
        add_paragraph(doc, line, "Normal")

    current_module = None
    case_index = 1
    for api in apis:
        if api.module_title != current_module:
            current_module = api.module_title
            add_paragraph(doc, current_module, "Heading 2")
        add_paragraph(doc, api.api_title, "Heading 3")
        build_case_table(doc, api, case_index)
        case_index += 1
        doc.add_paragraph()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    workspace = Path(r"D:\AI Demo\菁选")
    template = Path(r"D:/Temp/winzips/{B5E495D8-4CA9-43C2-B392-12EEE99F76E7}/用例示例.docx")
    markdown = workspace / "docs" / "接口测试用例.md"
    output = workspace / "docs" / "接口测试用例-Word版.docx"
    build_document(template, markdown, output)
    print(output)
