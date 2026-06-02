from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\AI Demo\菁选")
INPUT_MD = ROOT / "docs" / "测试计划.md"
OUTPUT_DOCX = ROOT / "docs" / "测试计划-Word版.docx"


TITLE = "测试计划"
SUBTITLE_PROJECT = "项目：菁选校园作品展示平台"
SUBTITLE_VERSION = "版本：1.0.0"
SUBTITLE_DATE = "日期：2026-05-24"


@dataclass
class HeadingBlock:
    level: int
    text: str


@dataclass
class ParagraphBlock:
    text: str


@dataclass
class ListBlock:
    ordered: bool
    items: list[str]


@dataclass
class TableBlock:
    rows: list[list[str]]


Block = HeadingBlock | ParagraphBlock | ListBlock | TableBlock


def normalize_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    return text.strip()


def parse_md(md_text: str) -> list[Block]:
    lines = md_text.splitlines()
    blocks: list[Block] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("> "):
            blocks.append(ParagraphBlock(normalize_inline(stripped[2:])))
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            blocks.append(HeadingBlock(len(m.group(1)), normalize_inline(m.group(2))))
            i += 1
            continue

        if stripped.startswith("|"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [normalize_inline(cell.strip()) for cell in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"-{3,}:?", c.replace(" ", "").replace(":", "")) for c in row):
                    rows.append(row)
                i += 1
            if rows:
                blocks.append(TableBlock(rows))
            continue

        if re.match(r"^[-*]\s+", stripped):
            items: list[str] = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(normalize_inline(re.sub(r"^\s*[-*]\s+", "", lines[i].strip())))
                i += 1
            blocks.append(ListBlock(False, items))
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(normalize_inline(re.sub(r"^\s*\d+\.\s+", "", lines[i].strip())))
                i += 1
            blocks.append(ListBlock(True, items))
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt == "---":
                break
            if nxt.startswith("> ") or nxt.startswith("|") or re.match(r"^(#{1,6})\s+", nxt):
                break
            if re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+\.\s+", nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        blocks.append(ParagraphBlock(normalize_inline(" ".join(paragraph_lines))))
    return blocks


def set_run_font(run, size: float, bold: bool = False, color: tuple[int, int, int] = (0, 0, 0)) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [("Heading 1", 20, (0, 0, 0)), ("Heading 2", 16, (0, 0, 0)), ("Heading 3", 14, (67, 67, 67))]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(20 if style_name == "Heading 1" else 18 if style_name == "Heading 2" else 16)
        style.paragraph_format.space_after = Pt(6 if style_name != "Heading 3" else 4)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(TITLE)
    set_run_font(run, 26)

    for text in (SUBTITLE_PROJECT, SUBTITLE_VERSION, SUBTITLE_DATE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, 12, color=(85, 85, 85))

    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["文档名称", "项目", "版本", "日期"]
    values = [TITLE, "菁选校园作品展示平台", "1.0.0", "2026-05-24"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for idx, text in enumerate(values):
        cell = table.cell(1, idx)
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_page_break()


def add_version_history(doc: Document) -> None:
    doc.add_heading("版本历史", level=1)
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["版本号", "生效日期", "版本说明/变更理由/变更内容", "作者", "备注"]
    values = ["1.0.0", "2026-05-24", "根据测试计划模板与现有测试计划 Markdown 生成首版 Word 文档", "Codex", "保留原测试计划正文内容"]
    for i, text in enumerate(headers):
        table.cell(0, i).text = text
    for i, text in enumerate(values):
        table.cell(1, i).text = text
    doc.add_page_break()


def add_toc_paragraph(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    text_run = paragraph.add_run("目录将在 Word 中自动更新。")
    set_run_font(text_run, 11, color=(85, 85, 85))
    run._r.append(fld_char_end)


def add_toc(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    p = doc.add_paragraph()
    add_toc_paragraph(p)
    doc.add_page_break()


def set_table_layout(table) -> None:
    table.autofit = False
    widths = []
    col_count = len(table.columns)
    if col_count == 2:
        widths = [Inches(1.8), Inches(4.7)]
    elif col_count == 3:
        widths = [Inches(1.4), Inches(1.6), Inches(3.5)]
    elif col_count == 4:
        widths = [Inches(0.9), Inches(1.6), Inches(2.6), Inches(1.4)]
    elif col_count == 5:
        widths = [Inches(0.9), Inches(1.1), Inches(2.6), Inches(0.9), Inches(1.0)]
    else:
        widths = [Inches(6.5 / max(col_count, 1))] * col_count

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = widths[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1.1
                if row_idx == 0:
                    for run in para.runs:
                        run.font.bold = True


def render_blocks(doc: Document, blocks: Iterable[Block]) -> None:
    for block in blocks:
        if isinstance(block, HeadingBlock):
            level = min(block.level, 3)
            doc.add_heading(block.text, level=level)
            continue

        if isinstance(block, ParagraphBlock):
            p = doc.add_paragraph(block.text)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            continue

        if isinstance(block, ListBlock):
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                p = doc.add_paragraph(item, style=style)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
            continue

        if isinstance(block, TableBlock):
            rows = block.rows
            cols = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r_idx, row in enumerate(rows):
                for c_idx, text in enumerate(row):
                    table.cell(r_idx, c_idx).text = text
            set_table_layout(table)
            doc.add_paragraph()


def build_doc(md_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    configure_styles(doc)
    add_cover(doc)
    add_version_history(doc)
    add_toc(doc)
    render_blocks(doc, parse_md(md_text))
    return doc


def update_toc_with_word(docx_path: Path) -> None:
    import win32com.client  # type: ignore

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path))
        try:
            for toc in doc.TablesOfContents:
                toc.Update()
            doc.Fields.Update()
            doc.Save()
        finally:
            doc.Close()
    finally:
        word.Quit()


def main() -> None:
    md_text = INPUT_MD.read_text(encoding="utf-8")
    doc = build_doc(md_text)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    try:
        update_toc_with_word(OUTPUT_DOCX)
    except Exception:
        pass
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
